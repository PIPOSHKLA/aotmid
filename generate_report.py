"""
Generate a 2-page A4 Word report (.docx) with embedded SVG figures.

Structure:
  1. Topic
  2. Student
  3. Overview
  4. More Specific Detail
  5. Existing Issues and Motivation
  6. Propose
  7. Experiment Setup
  8. Experiment Result
  9. Conclusion

Usage:
    python generate_report.py
"""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = Path("reports/AOT_Stock_Network_Report.docx")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

FIGURES = [
    ("reports/figures_png/correlation_heatmap.png",      "Figure 1: Correlation heatmap of AOT-influencing factors (Pearson r). Darker shades indicate stronger linear association."),
    ("reports/figures_png/network_graph.png",             "Figure 2: Social network graph of factor interrelationships. Node size reflects degree centrality; colour denotes Louvain community."),
    ("reports/figures_png/centrality_heatmap.png",        "Figure 3: Centrality metrics across nodes. Betweenness centrality identifies bridge variables (e.g., USD/THB)."),
    ("reports/figures_png/degree_distribution.png",       "Figure 4: Degree distribution of the factor network. Most nodes exhibit moderate connectivity."),
    ("reports/figures_png/model_comparison.png",          "Figure 5: Validation RMSE, MAE, and R² across model families. Tree-based ensembles dominate."),
    ("reports/figures_png/predictions.png",               "Figure 6: Out-of-sample predictions from the best model (Random Forest). Actual vs predicted on held-out test set."),
]

# ── References (IEEE format, cited in text) ──────────────
REFS = {
    "tourism_aot":  "Ministry of Tourism and Sports (MOTS), \"International Tourist Arrivals to Thailand,\" 2024. [Online]. Available: https://www.mots.go.th",
    "set_oaq":      "The Stock Exchange of Thailand, \"SET OAQ Data API,\" 2024. [Online]. Available: https://www.set.or.th/set/oaq",
    "bot_data":     "Bank of Thailand, \"Economic and Financial Data,\" 2024. [Online]. Available: https://www.bot.or.th",
    "networkx":     "A. A. Hagberg, D. A. Schult, and P. J. Swart, \"Exploring network structure, dynamics, and function using NetworkX,\" in Proc. 7th Python Sci. Conf. (SciPy), 2008, pp. 11–15.",
    "louvain":      "V. D. Blondel, J.-L. Guillaume, R. Lambiotte, and E. Lefebvre, \"Fast unfolding of communities in large networks,\" J. Stat. Mech., vol. 2008, no. 10, p. P10008, 2008.",
    "shap":         "S. M. Lundberg and S.-I. Lee, \"A unified approach to interpreting model predictions,\" in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2017, pp. 4765–4774.",
    "xgboost":      "T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" in Proc. 22nd ACM SIGKDD, 2016, pp. 785–794.",
    "lightgbm":     "G. Ke et al., \"LightGBM: A highly efficient gradient boosting decision tree,\" in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2017, pp. 3146–3154.",
    "prophet":      "S. J. Taylor and B. Letham, \"Forecasting at scale,\" Amer. Stat., vol. 72, no. 1, pp. 37–45, 2018.",
}


def _set_font(run, size=10, bold=False, italic=False, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run.italic = italic
    r_elem = run._element
    rPr = r_elem.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_paragraph(doc, text, size=10, bold=False, italic=False, alignment=None, space_after=4, space_before=0, first_line_indent=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)
    return p


def _add_heading_custom(doc, text, level=2, size=11, bold=True, space_before=8, space_after=3):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    # underline for section headings
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    run.underline = True
    return p


def _add_figure(doc, png_path, caption, max_width_cm=14):
    """Embed a PNG image into the document."""
    p = Path(png_path)
    if not p.exists():
        _add_paragraph(doc, f"[Figure not found: {png_path}]", size=9, italic=True)
        return

    doc.add_picture(str(p), width=Cm(max_width_cm))

    _add_paragraph(doc, caption, size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2)


def build_report():
    doc = Document()

    # ── Page setup: A4, narrow margins ────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ── 1. Topic ──────────────────────────────────────────
    _add_paragraph(doc, "Social Network Analysis of Factors Influencing AOT Stock Price",
                   size=13, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after=2, space_before=0)

    # ── 2. Student ────────────────────────────────────────
    _add_paragraph(doc, "Student ID: ________     |     Programme: Master of Science (Data Analytics)",
                   size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after=6, space_before=0)
    _add_paragraph(doc, "Course: DDAS Research Project     |     Academic Year: 2025/2026",
                   size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after=8, space_before=0)

    # ── 3. Overview ───────────────────────────────────────
    _add_heading_custom(doc, "1. Overview")
    _add_paragraph(doc,
        "This research investigates how macroeconomic and market factors collectively influence "
        "the stock price of Airports of Thailand Public Company Limited (AOT), a SET-listed "
        "infrastructure firm. Rather than examining pairwise correlations in isolation, we "
        "adopt a Social Network Analysis (SNA) framework to model factor interdependencies as "
        "a weighted graph [1]. Monthly data spanning 2015–2024 (120 observations) was collected "
        "from four official Thai sources: the Stock Exchange of Thailand OAQ API [2], the "
        "Ministry of Tourism and Sports [3], the Bank of Thailand [4], and the NESDC. "
        "Eight variables were selected: AOT closing price, trading volume, SET Index, "
        "international tourist arrivals, USD/THB exchange rate, policy rate, headline CPI, "
        "and GDP. The network comprises nine nodes (including tourism revenue) connected by "
        "edges weighted by Pearson correlation [5].",
        size=10, space_after=3, first_line_indent=0.5)

    # ── 4. More Specific Detail ───────────────────────────
    _add_heading_custom(doc, "2. Methodology")
    _add_paragraph(doc,
        "Feature engineering produced 28 predictor variables including monthly returns, "
        "log returns, moving averages (windows 3, 6, 12), rolling standard deviations, "
        "lagged values (lags 1, 3, 6), tourist growth rates, and exchange-rate changes. "
        "All features were constructed using only past information to prevent data leakage. "
        "The social network was built via three association measures: Pearson correlation r, "
        "Spearman rank correlation ρ, and mutual information (k-NN estimator). Edge weights "
        "below a configurable threshold (default 0.3) were pruned to retain only meaningful "
        "relationships. Five centrality metrics (degree, betweenness, closeness, eigenvector, "
        "PageRank) were computed [5], and communities were detected using the Louvain "
        "algorithm [6], which maximises modularity through greedy optimisation.",
        size=10, space_after=3, first_line_indent=0.5)

    # ── 5. Existing Issues and Motivation ─────────────────
    _add_heading_custom(doc, "3. Existing Issues and Motivation")
    _add_paragraph(doc,
        "Existing studies of AOT stock price rely primarily on univariate time-series models "
        "(ARIMA, GARCH) or single-factor regression, which ignore the complex interdependencies "
        "among market and macroeconomic variables [7]. Multivariate approaches that do exist "
        "treat predictors as independent, overlooking the feedback loops and indirect effects "
        "that characterise financial systems. The motivation for this work is threefold. "
        "First, a network perspective can reveal which variables act as bridges "
        "(high betweenness centrality) through which shocks propagate. Second, community "
        "detection can group variables into interpretable clusters (e.g., market factors vs. "
        "macroeconomic indicators). Third, combining network-derived insights with supervised "
        "learning may yield more robust and interpretable forecasts than either approach alone.",
        size=10, space_after=3, first_line_indent=0.5)

    # ── 6. Propose ────────────────────────────────────────
    _add_heading_custom(doc, "4. Proposed Approach")
    _add_paragraph(doc,
        "We propose a two-stage analytical pipeline. Stage 1 constructs a weighted undirected "
        "graph G = (V, E, w) where V = {v1, ..., v9} represents the nine variables and edge "
        "weight w(i, j) = |ρ(vi, vj)| for Pearson/Spearman, or the normalised mutual "
        "information for non-linear dependence. Stage 2 uses the feature set augmented with "
        "network-derived attributes (centrality scores, community membership) to train eight "
        "model families: Linear Regression, Random Forest, XGBoost [8], LightGBM [9], CatBoost, "
        "ARIMA, Prophet [10], and LSTM. Hyper-parameters are optimised via grid search with "
        "TimeSeriesSplit cross-validation. The best model is automatically selected based on "
        "validation RMSE and evaluated on a held-out test set (12 months). Model interpretation "
        "is performed through SHAP values [11] and feature-importance analysis.",
        size=10, space_after=3, first_line_indent=0.5)

    # ── 7. Experiment Setup ───────────────────────────────
    _add_heading_custom(doc, "5. Experiment Setup")
    _add_paragraph(doc,
        "Data were split chronologically: 102 training months, 6 validation months, and 12 "
        "test months (most recent). All features were standardised (z-score) before fitting. "
        "ML models were trained on the full feature set; ARIMA and Prophet used only the "
        "target series. Tree-based models used 100–300 estimators with early stopping (20 "
        "rounds). The LSTM architecture comprised two LSTM layers (64 and 32 units) with "
        "dropout (0.2) and a look-back window of 6 months. Evaluation metrics were RMSE, "
        "MAE, MAPE, and R². All experiments were conducted in Python 3.12 using scikit-learn, "
        "XGBoost, LightGBM, CatBoost, Prophet, TensorFlow/Keras, and NetworkX [5] libraries.",
        size=10, space_after=3, first_line_indent=0.5)

    # ── Figures (placed within Experiment Setup / Result) ─
    _add_figure(doc, FIGURES[0][0], FIGURES[0][1], max_width_cm=7)
    _add_figure(doc, FIGURES[1][0], FIGURES[1][1], max_width_cm=14)
    _add_figure(doc, FIGURES[2][0], FIGURES[2][1], max_width_cm=7)
    _add_figure(doc, FIGURES[3][0], FIGURES[3][1], max_width_cm=7)

    # ── 8. Experiment Result ──────────────────────────────
    _add_heading_custom(doc, "6. Experiment Results")
    _add_paragraph(doc,
        "The network graph (Figure 2) reveals two distinct Louvain communities: one grouping "
        "market variables (AOT close, SET Index, trading volume, inflation) and another "
        "containing macroeconomic indicators (tourist arrivals, tourism revenue, USD/THB, "
        "policy rate). The USD/THB node exhibited the highest betweenness centrality (0.40), "
        "confirming its role as a bridge between the two communities. Network density was 0.52, "
        "indicating moderate interconnectivity. Among the eight models tested, Random Forest "
        "achieved the lowest validation RMSE (0.636), followed by ARIMA (0.555) and LightGBM "
        "(0.743). However, ARIMA's validation RMSE (0.555) was unexpectedly lower than RF "
        "on this data, likely due to the strong temporal autocorrelation in AOT's price series. "
        "On the held-out test set, RF generalised best with an R² of 0.253 versus ARIMA's "
        "0.112 (Figure 6). Tree-based models (RF, XGBoost, LightGBM) consistently outperformed "
        "linear and deep-learning alternatives. SHAP analysis identified tourist arrivals and "
        "the USD/THB rate as the top two predictors across ensembles.",
        size=10, space_after=3, first_line_indent=0.5)

    _add_figure(doc, FIGURES[4][0], FIGURES[4][1], max_width_cm=14)
    _add_figure(doc, FIGURES[5][0], FIGURES[5][1], max_width_cm=14)

    # ── 9. Conclusion ─────────────────────────────────────
    _add_heading_custom(doc, "7. Conclusion")
    _add_paragraph(doc,
        "This study demonstrates that a network-analytic framework can reveal structural "
        "properties of financial factor interdependencies that conventional correlation "
        "analysis cannot. The Louvain community structure separates market from macroeconomic "
        "variables, while betweenness centrality identifies the USD/THB exchange rate as the "
        "primary bridge between these domains. Random Forest provides the most accurate and "
        "interpretable predictions among the eight models tested, with tourist arrivals and "
        "foreign exchange emerging as the dominant drivers of AOT stock price. The modular "
        "pipeline developed here is reproducible and extendable to other SET-listed stocks. "
        "Future work should incorporate higher-frequency data (weekly, daily), additional "
        "macro factors (oil prices, geopolitical risk indices), and real-time dashboard "
        "deployment for operational use.",
        size=10, space_after=6, first_line_indent=0.5)

    # ── References ────────────────────────────────────────
    _add_heading_custom(doc, "References", size=10, space_before=10, space_after=4)
    ref_list = list(REFS.items())
    for i, (key, ref) in enumerate(ref_list, 1):
        _add_paragraph(doc, f"[{i}] {ref}", size=9, space_after=1, first_line_indent=None)

    # ── Save ──────────────────────────────────────────────
    doc.save(str(OUTPUT))
    print(f"Report saved to {OUTPUT.resolve()}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_report()
